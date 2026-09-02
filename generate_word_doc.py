import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def create_document():
    doc = docx.Document()

    # Set page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    title_style = doc.styles.add_style('DocTitle', docx.enum.style.WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Calibri'
    title_font.size = Pt(22)
    title_font.bold = True
    title_font.color.rgb = RGBColor(15, 23, 42) # Slate 900

    # 1. HEADER & TITLE
    title_p = doc.add_paragraph("GD/H126 INTELLEKTUAL CPR VA BEMOR SIMULYATORI PLATFORMASI", style='DocTitle')
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub_p = doc.add_paragraph("Tibbiy Ta'lim, Datchiklar Integratsiyasi, OSCE Imtihon Tizimi va Video Taqdimot Qo'llanmasi")
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_p.runs[0].font.size = Pt(12)
    sub_p.runs[0].font.color.rgb = RGBColor(79, 70, 229) # Indigo 600
    sub_p.runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Info Box Table
    info_table = doc.add_table(rows=3, cols=2)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    info_data = [
        ("Loyiha Nomi:", "MedLife: AI Bemor Simulyatori, Vital Monitor va Yurak-O'pka Reanimatsiyasi Tizimi"),
        ("Asosiy Manzillar:", "Web: https://med-bemor-simulyator.onrender.com/ | Portlar: 8000, 8500, 8600"),
        ("Tuzuvchilar:", "1-Ishtirokchi (Dasturiy Ta'minot) & 2-Ishtirokchi (Hardware / Datchiklar)")
    ]
    for idx, (label, val) in enumerate(info_data):
        row = info_table.rows[idx]
        cell_lbl, cell_val = row.cells[0], row.cells[1]
        cell_lbl.width = Inches(2.0)
        cell_val.width = Inches(4.8)
        
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.add_run(label).bold = True
        p_lbl.runs[0].font.size = Pt(10)
        p_lbl.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        
        p_val = cell_val.paragraphs[0]
        p_val.add_run(val)
        p_val.runs[0].font.size = Pt(10)
        p_val.runs[0].font.color.rgb = RGBColor(15, 23, 42)

        set_cell_background(cell_lbl, "F1F5F9")
        set_cell_background(cell_val, "F8FAFC")
        set_cell_margins(cell_lbl, 80, 80, 120, 120)
        set_cell_margins(cell_val, 80, 80, 120, 120)

    doc.add_paragraph().paragraph_format.space_after = Pt(15)

    # ==================== 1-BO'LIM ====================
    h1 = doc.add_heading("1-BO'LIM. 2 KISHILIK RASMIY VIDEO TAQDIMOT SSENARIYSI", level=1)
    h1.runs[0].font.color.rgb = RGBColor(30, 58, 138) # Blue 900
    h1.runs[0].font.size = Pt(15)

    intro_sc = doc.add_paragraph("Ushbu ssenariy video tasvirga olish jarayonida 2 kishi (Dasturchi va Hardware muhandisi) birgalikda, ixcham va dinamik tarzda loyihani tushuntirishi uchun tayyorlangan. Har bir jumla aytilganda ekranda unga mos video lavha (B-roll) ko'rinadi.")
    intro_sc.runs[0].font.italic = True
    intro_sc.runs[0].font.size = Pt(10.5)

    script_parts = [
        ("1. KIRISH VA TANISHTIRUV (~20 soniya)", [
            ("👤 1-Ishtirokchi (Dasturchi - Siz):", "Assalomu alaykum! Men [Ismingiz] — ushbu loyihaning dasturiy ta'minot muhandisiman."),
            ("👤 2-Ishtirokchi (Muhandis - Sherigingiz):", "Men esa [Sherigingiz ismi] — loyihaning apparat qismi, ya'ni maniken va sensorlar bo'yicha muhandisiman."),
            ("👤 1-Ishtirokchi (Siz):", "Bugun sizlarga tibbiy ta'limni yangi bosqichga olib chiquvchi innovatsion platformamiz — MedLife tibbiy simulyatorlar majmuasini taqdim etamiz.")
        ]),
        ("2. APPARAT VA DATCHIKLAR QISMI (~40 soniya)", [
            ("📹 Kadrda / Ekranda:", "[Maniken, ESP32 mikrokontrolleri, datchiklar va simlar montaji ko'rsatiladi]"),
            ("👤 2-Ishtirokchi (Sherigingiz):", "Loyihamizning apparat qismida biz GD/H126 tibbiy manikenini to'liq modernizatsiya qildik:\n"
                                               "• Ko'krak qafasiga massaj kuchini 0 dan 60 kg gacha o'lchovchi datchik;\n"
                                               "• Nafas yo'llari va oshqozonga yuqori aniqlikdagi pnevmatik bosim datchiklari;\n"
                                               "• To'sh suyagi va qo'l tomiriga kontaktli sensorlar va pulsator o'rnatdik.\n"
                                               "Barcha datchiklardan kelayotgan ma'lumotlar ESP32 mikrokontrolleri orqali qayta ishlanib, 0 millisekundlik kechikish bilan dasturga uzatiladi. Dasturiy qismni esa sherigim tushuntiradi.")
        ]),
        ("3. DASTURIY MODULLAR (~30 soniya)", [
            ("📹 Kadrda / Ekranda:", "[AI Bemor, ICU Vital Monitor va Yurak-O'pka Reanimatsiyasi simulyatori navbatma-navbat ko'rsatiladi]"),
            ("👤 1-Ishtirokchi (Siz):", "Dasturimiz 3 ta asosiy moduldan iborat:\n"
                                       "1. MedLife: AI Bemor Simulyatori — talaba 14 xil kasallik bo'yicha bemor bilan o'zbek tilida jonli muloqot qiladi;\n"
                                       "2. Vital Monitor — bemorning EKG, puls, SpO2 va bosimini real vaqtda jonli ko'rsatadi;\n"
                                       "3. Yurak-O'pka Reanimatsiyasi Simulyatori — datchiklar bilan integratsiyalashgan to'liq raqamli CPR boshqaruv va imtihon tizimi.")
        ]),
        ("4. TEZKOR IMTIHON VA KLINIK FUNKSIYALAR (~40 soniya)", [
            ("📹 Kadrda / Ekranda:", "[Teskari hisob, datchiklar ko'rsatkichlari, Bemorning tirilishi va A4 Protokol ko'rsatiladi]"),
            ("👤 1-Ishtirokchi (Siz):", "Imtihon jarayoni juda qulay tashkil qilingan:\n"
                                       "• Talabani ro'yxatga olish va 3 soniyalik ovozli teskari hisob (3... 2... 1... START!);\n"
                                       "• Xolis klinik baholash: Kompressiya chuqurligi (38-55kg), recoil, nafas me'yori (2.0-3.0 kPa) va oshqozon xatolarini avtomatik hisoblaydi;\n"
                                       "• Bemorning Tirilishi (ROSC effekti): Talaba a'lo topshirsa, maniken ko'kragi mustaqil nafas oladi, ko'zlari ochiladi va ovozli minnatdorchilik bildiradi;\n"
                                       "• A4 Protokol: Imtihon tugashi bilan birgina tugma orqali rasmiy muhrli baholash varaqasini chop etish mumkin;\n"
                                       "• Jurnal va Qidiruv: Barcha talabalar natijalari xotirada saqlanadi, saralanadi va qidiriladi.")
        ]),
        ("5. YAKUNIY XULOSA (~15 soniya)", [
            ("👤 2-Ishtirokchi (Sherigingiz):", "Ushbu tizim talabalarga real klinik vaziyatni 100% his qilish imkonini beradi."),
            ("👤 1-Ishtirokchi (Siz):", "O'qituvchilar uchun esa imtihonlarni shaffof va oson o'tkazishni ta'minlaydi. E'tiboringiz uchun rahmat!")
        ])
    ]

    for sec_title, lines in script_parts:
        h2 = doc.add_heading(sec_title, level=2)
        h2.runs[0].font.color.rgb = RGBColor(14, 116, 144) # Cyan 700
        h2.runs[0].font.size = Pt(12)

        for speaker, text in lines:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.2)
            p.paragraph_format.space_after = Pt(4)
            
            run_spk = p.add_run(speaker + " ")
            run_spk.bold = True
            if "1-Ishtirokchi" in speaker:
                run_spk.font.color.rgb = RGBColor(37, 99, 235) # Blue
            elif "2-Ishtirokchi" in speaker:
                run_spk.font.color.rgb = RGBColor(16, 185, 129) # Green
            else:
                run_spk.font.color.rgb = RGBColor(168, 85, 247) # Purple

            run_txt = p.add_run(text)
            run_txt.font.size = Pt(10.5)

    doc.add_page_break()

    # ==================== 2-BO'LIM ====================
    h1_2 = doc.add_heading("2-BO'LIM. TIZIMNING TEXNIK ARXITEKTURASI VA DATCHIKLAR", level=1)
    h1_2.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h1_2.runs[0].font.size = Pt(15)

    doc.add_paragraph("Platforma apparat (Hardware) va dasturiy (Software) qismlarining mukammal sinxronizatsiyasi asosida ishlaydi:")

    hw_points = [
        ("Ko'krak Massaj Kuchi Sensor (Strain Gauge):", "0 dan 60 kg gacha bosimni o'lchaydi. 38 - 55 kg oralig'i to'g'ri (Norma) deb baholanadi. <38 kg (Sayoz) va >55 kg (Haddan tashqari qattiq) xatolar hisoblanadi."),
        ("Pnevmatik O'pka Nafas Sensori (Airway Pressure):", "0 dan 4.0 kPa gacha ventilyatsiya bosimini qayd etadi. 2.0 - 3.0 kPa (20 - 30 cmH2O) Ambu qopi bilan nafas berishning xavfsiz klinik me'yori hisoblanadi."),
        ("Oshqozon Bosimi Sensori (Stomach Pressure):", "Noto'g'ri nafas berilganda havoning oshqozonga ketishini nazorat qiladi. Oshqozonda bosim >0.8 kPa bo'lganda ogohlantirish beradi."),
        ("To'sh Suyagi Kontakt Sensori (PIN 13):", "Qo'lning to'g'ri anatomiya nuqtasiga (to'sh suyagi o'rtasiga) qo'yilganligini tekshiradi."),
        ("O'ng Qo'l Inyeksiya Sensori (PIN 4):", "Ukol / inyeksiya qilinganda dori yuborilganini fiksatsiya qiladi."),
        ("ESP32 UART Telemetriya Ko'prigi:", "115200 bod tezlikda barcha telemetriyani JSON formatida portlar (8600, 8500, 8000) orqali WebSocket oqimiga uzatadi (Kechikish: 0ms).")
    ]

    for title, desc in hw_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(title + " ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ==================== 3-BO'LIM ====================
    h1_3 = doc.add_heading("3-BO'LIM. DASTURIY IMKONIYATLAR VA KLINIK MEZONLAR", level=1)
    h1_3.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h1_3.runs[0].font.size = Pt(15)

    sw_points = [
        ("Erkin Mashq va Imtihon Rejimlari:", "Katta, qulay tugmalar orqali tanlanadi. Erkin mashqda talaba cheklovlarsiz mashq qiladi, barcha datchiklar real vaqtda ishlab turadi. Imtihon rejimida esa 2 daqiqalik rasmiy vaqt ishlaydi."),
        ("3 Soniyalik Ovozli Teskari Hisob:", "Imtihon boshlanganda 3... 2... 1... START ovozli signali chalinib, talabaga qo'llarini joylashtirib tayyor bo'lish imkonini beradi."),
        ("Bemorning Tirilishi (ROSC Effekti):", "Imtihon muvaffaqiyatli topshirilganda (≥80% ball va to'g'ri CPR), maniken ko'kragi mustaqil 3.6 soniyalik ritmda nafas oladi, ko'zlari ochiladi, 75 BPM sinus puls tiklanadi, konfetti otiladi va bemor o'zbek tilida minnatdorchilik bildiradi."),
        ("A4 Rasmiy Baholash Protokoli va Chop Etish:", "Imtihon yakunida rasmiy muhrlangan, talaba va o'qituvchi imzolari qo'yiladigan A4 protokoli printerga chiqariladi yoki PDF qilib saqlanadi."),
        ("Doimiy Imtihon Jurnali (Xotira):", "Barcha talabalar natijalari brauzer xotirasida doimiy saqlanadi. Natijalarni qidirish, o'tgan/yiqilgan bo'yicha filtrlash va saralash funksiyalari mavjud.")
    ]

    for title, desc in sw_points:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r_t = p.add_run(title + " ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(15, 23, 42)
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # ==================== 4-BO'LIM ====================
    h1_4 = doc.add_heading("4-BO'LIM. HAVOLALAR VA ISHGA TUSHIRISH YO'RIQNOMASI", level=1)
    h1_4.runs[0].font.color.rgb = RGBColor(30, 58, 138)
    h1_4.runs[0].font.size = Pt(15)

    doc.add_paragraph("Platformadan foydalanish uchun quyidagi havolalar mavjud:")

    links = [
        ("Asosiy Bemor Simulyatori:", "https://med-bemor-simulyator.onrender.com/"),
        ("ICU Vital Monitor:", "https://med-bemor-simulyator.onrender.com/monitor"),
        ("GD/H126 CPR Imtihon Pulti:", "https://med-bemor-simulyator.onrender.com/console (yoki /pult, /exam)"),
        ("Lokal kompyuterda ishga tushirish:", "Papkadagi '🎮 Manikin_Pult_Test.bat' faylini ikki marta bosing.")
    ]

    for title, desc in links:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        r_t = p.add_run(title + " ")
        r_t.bold = True
        r_d = p.add_run(desc)
        r_d.font.size = Pt(10.5)

    filename = "GD_H126_Bemor_Maniken_Loyiha_va_Video_Ssenariy.docx"
    doc.save(filename)
    print(f"Successfully generated: {filename}")

if __name__ == "__main__":
    create_document()
