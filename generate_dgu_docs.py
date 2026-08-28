import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>'))

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_code_block(doc, filename, code_text):
    h = doc.add_heading(filename, level=2)
    h.runs[0].font.name = 'Consolas'
    h.runs[0].font.size = Pt(12)
    h.runs[0].font.bold = True
    h.runs[0].font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(30, 41, 59)

def generate_dgu_1():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Titul Varag'i
    p1 = doc.add_paragraph("Deponentlanadigan materiallarning titul varag‘i")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].font.size = Pt(14)
    p1.runs[0].font.bold = True

    p2 = doc.add_paragraph("EHM uchun dastur (Ma’lumotlar bazasi) nomi:\n“Gemini Live Sun'iy Intellekt Asosidagi Interaktiv Virtual Bemor Simulyatori va Klinik Muloqot Dasturi”")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.runs[0].font.size = Pt(12)
    p2.runs[0].font.bold = True

    doc.add_paragraph("Huquq ega(lar)si:\n1. ....................................................................................................")
    doc.add_paragraph("Muallif(lar):\n1. ....................................................................................................\n2. ....................................................................................................")

    p_div = doc.add_paragraph("EHM uchun dasturni identifikatsiya qiluvchi materiallar dastlabki matni (Dastur kodi)")
    p_div.runs[0].font.size = Pt(11)
    p_div.runs[0].font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Read files
    try:
        with open("bemor_simulyator.py", "r", encoding="utf-8") as f:
            code_bemor = f.read()
    except Exception:
        code_bemor = "# bemor_simulyator.py"

    try:
        with open("web_app.py", "r", encoding="utf-8") as f:
            code_web = f.read()
    except Exception:
        code_web = "# web_app.py"

    add_code_block(doc, "bemor_simulyator.py", code_bemor)
    add_code_block(doc, "web_app.py", code_web)

    filename = "DGU_1_AI_Bemor_Simulyatori_Dastur_Kodi.docx"
    doc.save(filename)
    print(f"Generated: {filename}")

def generate_dgu_2():
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Titul Varag'i
    p1 = doc.add_paragraph("Deponentlanadigan materiallarning titul varag‘i")
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.runs[0].font.size = Pt(14)
    p1.runs[0].font.bold = True

    p2 = doc.add_paragraph("EHM uchun dastur (Ma’lumotlar bazasi) nomi:\n“GD/H126 Reanimatsiya Vital Monitori, ESP32 Telemetriyali Maniken Pulti va CPR OSCE Imtihon Baholash Dasturi”")
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.runs[0].font.size = Pt(12)
    p2.runs[0].font.bold = True

    doc.add_paragraph("Huquq ega(lar)si:\n1. ....................................................................................................")
    doc.add_paragraph("Muallif(lar):\n1. ....................................................................................................\n2. ....................................................................................................")

    p_div = doc.add_paragraph("EHM uchun dasturni identifikatsiya qiluvchi materiallar dastlabki matni (Dastur kodi)")
    p_div.runs[0].font.size = Pt(11)
    p_div.runs[0].font.bold = True
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # Read files
    try:
        with open("vital_monitor.py", "r", encoding="utf-8") as f:
            code_vital = f.read()
    except Exception:
        code_vital = "# vital_monitor.py"

    try:
        with open("manikin_console.py", "r", encoding="utf-8") as f:
            code_manikin = f.read()
    except Exception:
        code_manikin = "# manikin_console.py"

    try:
        with open("esp32_serial_bridge.py", "r", encoding="utf-8") as f:
            code_bridge = f.read()
    except Exception:
        code_bridge = "# esp32_serial_bridge.py"

    add_code_block(doc, "vital_monitor.py", code_vital)
    add_code_block(doc, "manikin_console.py", code_manikin)
    add_code_block(doc, "esp32_serial_bridge.py", code_bridge)

    filename = "DGU_2_Vital_Monitor_va_CPR_Imtihon_Dastur_Kodi.docx"
    doc.save(filename)
    print(f"Generated: {filename}")

if __name__ == "__main__":
    generate_dgu_1()
    generate_dgu_2()
